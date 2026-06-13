import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

#regenerates docs/StratLab_Documentation.docx from scratch so the document always
#reflects the current state of the app. run `python make_docs.py` after structural changes.

#colours mirror the app palette in theme.py so the document looks on-brand
PRIMARY = RGBColor(0x43, 0x38, 0xCA)
MUTED = RGBColor(0x64, 0x74, 0x8B)
CODE_FILL = "F1F5F9"

OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "docs", "StratLab_Documentation.docx")


#--- small formatting helpers so the content below stays readable ---

def add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = PRIMARY


def add_heading(doc: Document, text: str, size: int = 15) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = PRIMARY


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_code(doc: Document, code_text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    #light grey background so code blocks stand out from prose
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), CODE_FILL)
    properties.append(shading)


#--- the document content, section by section ---

def write_intro(doc: Document) -> None:
    add_title(doc, "StratLab: Project Documentation")
    add_para(doc, "StratLab is a web app for building, backtesting and comparing simple "
                  "investment strategies on real historical data. It works like a flight "
                  "simulator for investing: you can test a strategy on past data before "
                  "risking real money.")
    add_para(doc, "It is built with Python and Streamlit. All the finance logic, meaning the "
                  "data loading, strategies, backtest engine and every metric, lives in "
                  "core/, which knows nothing about Streamlit. The Streamlit pages just call "
                  "into core/ and show the result.")

    add_heading(doc, "About the authors")
    add_para(doc, "StratLab was built by two Master in Banking and Finance students at the "
                  "University of St. Gallen, as a tool to backtest simple strategies in "
                  "minutes instead of hours, without writing code or paying for data.")

    add_heading(doc, "AI Usage Declaration")
    add_para(doc, "Claude was used to clean up and structure the code. ChatGPT was used to "
                  "understand the Yahoo Finance API and the strategies conceptually. The logic "
                  "and design decisions are our own.")
    add_caption(doc, "Additional source: https://ranaroussi.github.io/yfinance/")

    add_heading(doc, "Sources for the finance formulas")
    add_para(doc, "The metrics and formulas in core/measures.py (Sharpe, Sortino, Treynor, VaR, "
                  "CAPM/Fama-French factor models, the Markowitz tangency portfolio, EWMA "
                  "covariance, etc.) follow the conventions taught in the AQAM (Advanced "
                  "Quantitative Asset Management) course at the University of St. Gallen, plus "
                  "the academic sources cited in the Glossary. They are not AI-generated.")

    add_heading(doc, "Tech stack")
    add_bullets(doc, [
        "Streamlit: the web interface and multi-page navigation",
        "pandas / numpy: price data and strategy maths",
        "yfinance: historical prices from Yahoo Finance",
        "scipy: skewness, kurtosis, the Jarque-Bera test, normal quantiles",
        "scikit-learn: logistic regression for the next-day forecast feature",
        "matplotlib: charts",
        "Supabase (Postgres): stores users and saved strategies",
        "python-docx: generates this document",
    ])

    add_heading(doc, "Folder structure")
    add_code(doc,
        "stratlab_final/\n"
        "    core/                  <- finance engine (pure Python, no Streamlit)\n"
        "        data.py            <- PriceLoader, Fama-French download\n"
        "        strategy.py        <- the five strategies\n"
        "        backtest.py        <- Backtest.run()\n"
        "        measures.py        <- finance formulas + Analytics class\n"
        "        forecast.py        <- next-day direction forecast (logistic regression)\n"
        "    services/              <- persistence (Supabase)\n"
        "        db.py, auth.py, storage.py\n"
        "    pages/                 <- Streamlit pages\n"
        "        0_Home.py, 1_Builder.py, 2_Evaluate.py,\n"
        "        3_Compare.py, 4_Explore.py, 5_Glossary.py\n"
        "    app.py                 <- entry point: login + navigation\n"
        "    theme.py               <- design tokens (colours, font, spacing)\n"
        "    ui.py                  <- reusable UI components + shared charts\n"
        "    strategies_config.py   <- strategy registry for the UI\n"
        "    prices.py              <- cached price/fundamentals loader\n"
        "    tickers.py             <- searchable universe of stocks/ETFs\n"
        "    assets/                <- logos, favicon, style.css\n"
        "    docs/                  <- this document\n"
        "    make_docs.py           <- regenerates this document\n")
    add_para(doc, "The dependencies only go one way: pages/ use services/, prices.py and ui.py, "
                  "which in turn use core/. Nothing in core/ imports Streamlit or services/, so "
                  "the finance engine can be tested or reused on its own.")


def write_core(doc: Document) -> None:
    add_heading(doc, "The core/ engine")
    add_para(doc, "Everything that computes a number lives in core/. The UI never does finance "
                  "maths itself.")

    add_heading(doc, "core/data.py: PriceLoader", size=13)
    add_bullets(doc, [
        "PriceLoader().get_prices(tickers, start, end) returns a DataFrame of daily "
        "adjusted-close prices from Yahoo Finance, one column per ticker. A ticker that "
        "fails to download is skipped silently.",
        "load_fama_french_factors(model) downloads monthly Fama-French factors "
        "(\"FF3\" or \"FF5\") from Ken French's data library, wrapped in try/except so a "
        "failed download returns an empty DataFrame instead of crashing.",
        "core/data.py is deliberately cache-free; prices.py adds an st.cache_data layer for "
        "the UI.",
    ])
    add_para(doc, "Error handling: _download_from_yahoo and load_fama_french_factors each "
                  "catch any exception during download or parsing and return an empty "
                  "Series/DataFrame, so one bad ticker or a Ken French outage does not crash "
                  "the page.")

    add_heading(doc, "core/strategy.py: the five strategies", size=13)
    add_para(doc, "Every strategy is a small class with the same interface. __init__ takes its "
                  "own parameters plus a shared rebalance_every value (the number of trading "
                  "days between rebalances, default 252). get_weights(prices) returns a dict "
                  "of target weights for \"today\", or the same weights as before if it is not "
                  "yet time to rebalance. Three shared helper functions, _should_rebalance(...), "
                  "_normalise_weights(...) and build_weights(...), avoid repeating this logic "
                  "in every strategy.")
    add_bullets(doc, [
        "EqualWeight: splits the portfolio evenly across all tickers (the \"1/N\" rule). "
        "DeMiguel, Garlappi and Uppal (2009) showed this is a tough benchmark to beat.",
        "FixedWeight: uses weights chosen by the user (for example 60/40 stocks/bonds).",
        "Momentum: ranks assets by their past return over a lookback window and goes long "
        "the top performers (optionally short the worst ones), based on Jegadeesh and "
        "Titman (1993).",
        "InverseVolatility: weights each asset by the inverse of its volatility, so calmer "
        "assets get a bigger weight.",
        "MeanVariance: the Markowitz (1952) tangency portfolio. Weights are w = Sigma^-1 * "
        "mu, normalised so they sum to 1, with negative weights clipped to zero. It can "
        "optionally use an exponentially weighted (EWMA / RiskMetrics) covariance matrix "
        "instead of the plain sample covariance.",
    ])
    add_para(doc, "MeanVariance and InverseVolatility need a minimum amount of return history "
                  "before their covariance or volatility is defined (more observations than "
                  "assets, or at least two observations). Until then they hold cash (empty "
                  "weights), so the equity curve stays flat instead of showing NaN.")
    add_para(doc, "Error handling: _normalise_weights returns all-zero weights instead of "
                  "dividing by zero (for example, when MeanVariance clips every weight to "
                  "zero). EqualWeight and Momentum stay in cash if there are no tickers or top "
                  "performers, Momentum treats a 0 starting price as a 0% return, "
                  "InverseVolatility gives a zero-volatility ticker a weight of 0, and "
                  "MeanVariance uses np.linalg.pinv instead of np.linalg.inv so a singular "
                  "covariance matrix (for example, two identical tickers) does not raise an "
                  "error.")

    add_heading(doc, "core/backtest.py: Backtest", size=13)
    add_para(doc, "Backtest().run(strategy, prices) walks forward through the price history "
                  "one day at a time. On each day it gives strategy.get_weights() only the "
                  "prices up to that day, so there is no look-ahead bias. It then works out "
                  "the weighted daily return from each ticker's price change, and grows the "
                  "portfolio value (starting at 1.0) by that return. The result is a Series, "
                  "stored on self.portfolio and also returned. All KPIs are then computed from "
                  "this series using core/measures.py; Backtest itself has no KPI methods.")
    add_para(doc, "Error handling: a ticker missing from the price data for a given day is "
                  "skipped (it contributes nothing to that day's return), and a ticker with a "
                  "0 price the day before is also skipped for that day, instead of dividing by "
                  "zero. If given 0 or 1 rows of prices (a 0-length period), the day-by-day "
                  "loop never runs and run() returns an empty Series rather than raising an "
                  "error; the Evaluate and Compare pages avoid this case by checking that there "
                  "are at least 2 rows of price data before calling Backtest at all. As an extra "
                  "safety net, each day's calculation is also wrapped in its own try/except: if "
                  "a strategy raises an unexpected error on one day, that day simply carries the "
                  "portfolio value forward unchanged instead of crashing the whole backtest.")

    add_heading(doc, "core/measures.py: finance formulas and the Analytics class", size=13)
    add_para(doc, "A library of stand-alone functions, each with a short docstring explaining "
                  "the formula:")
    add_bullets(doc, [
        "Returns: simple and log returns, average/annualised return, CAGR, volatility "
        "(daily, monthly, annualised).",
        "Risk-adjusted: Sharpe, Sortino and Treynor ratios (risk-free-adjusted).",
        "Distribution: skewness, excess kurtosis (K-3), Jarque-Bera normality test.",
        "Risk: maximum drawdown, historic VaR (empirical quantile) and "
        "variance-covariance VaR (t*mu + sqrt(t)*sigma*z, for 1-day and 1-month horizons).",
        "Factor models: CAPM, Fama-French 3- and 5-factor monthly regressions "
        "(alpha, beta, t-/p-values, R-squared, appraisal ratio).",
        "Portfolio construction: EWMA covariance (RiskMetrics) for the MeanVariance "
        "strategy, common_window to align tickers that start on different dates, and "
        "the cash/risk-free overlay, trading/management cost application and annual "
        "turnover used to turn a gross backtest into a net one.",
    ])
    add_para(doc, "factor_model(asset_excess, factors, periods_per_year) and its one-factor "
                  "wrapper capm(...) regress an asset's excess return on one or more factors "
                  "(R - Rf = alpha + B * factors + eps) using the private helper _ols. _ols "
                  "performs ordinary least squares with an intercept: it builds a design "
                  "matrix X = [1, factors], solves for the coefficients with "
                  "np.linalg.lstsq, then computes the residuals as y - X @ coef (matrix "
                  "multiplication: the fitted values), the sum of squared residuals as "
                  "residuals @ residuals (a dot product), and R-squared. The coefficient "
                  "standard errors (for the t-stats and p-values) come from "
                  "sigma^2 * (X'X)^-1, using np.linalg.pinv instead of an exact inverse so "
                  "near-collinear factors do not raise an error.")
    add_para(doc, "Analytics(values, benchmark_values, rf_annual) bundles these into one object "
                  "so a page can pull every measure for one portfolio (and its benchmark) "
                  "without re-deriving returns each time. This is what Evaluate and Compare "
                  "use for their KPI cards and tables.")
    add_para(doc, "log_returns is not currently called from any page or other core function. "
                  "It was added during backend development for possible future use (e.g. "
                  "log-return-based statistics) and left in for that reason.")
    add_para(doc, "Error handling: sharpe_ratio, sortino_ratio, treynor_ratio, cagr and "
                  "historic_var return NaN instead of dividing by zero or indexing into empty "
                  "data, for very short or unusual return series. The factor-model helper _ols "
                  "returns NaN for statistics that need more observations than parameters, and "
                  "uses pinv instead of inv so near-collinear factors do not raise an error. "
                  "Analytics.beta returns NaN instead of dividing by zero if the benchmark has "
                  "zero (or undefined) variance, for example a flat-priced benchmark or too few "
                  "aligned observations.")

    add_heading(doc, "core/forecast.py: next-day direction forecast", size=13)
    add_para(doc, "predict_direction(prices, n_lags=10) is a small machine-learning add-on "
                  "used on the Explore and Evaluate pages. Each day's features are the "
                  "previous n_lags daily returns, and the label is whether that day's return "
                  "was positive. A scikit-learn LogisticRegression model is trained on a "
                  "chronological 80/20 split, so the test set is always later in time than "
                  "the training set. The function returns the test-set accuracy plus the "
                  "predicted probability that tomorrow's return will be positive, based on "
                  "the most recent n_lags returns. It returns None if there is not enough "
                  "price history (fewer than n_lags + 30 return observations) to train and "
                  "test the model.")
    add_para(doc, "Error handling: predict_direction returns None if there is not enough "
                  "price history to train and test a model. The Explore and Evaluate pages "
                  "show an info message instead of the Forecast section in that case.")


def write_services(doc: Document) -> None:
    add_heading(doc, "services/: persistence layer")
    add_para(doc, "Users and saved strategies live in a hosted Supabase (Postgres) database, "
                  "reached through a single cached client. Credentials sit in "
                  ".streamlit/secrets.toml (excluded from version control).")
    add_bullets(doc, [
        "db.py: get_client() builds the Supabase client from "
        ".streamlit/secrets.toml (the [supabase] url and key) and is decorated with "
        "@st.cache_resource, so it is created once and reused across reruns instead of "
        "opening a new connection on every interaction. run_query(builder) takes a "
        "Supabase query builder (e.g. client.table(\"users\").select(\"password\").eq(...)), "
        "calls builder.execute().data, and returns its data, or shows st.error(...) and "
        "returns None if anything goes wrong, so a database error shows up as a message "
        "instead of crashing the app. Every function in auth.py and storage.py goes "
        "through this one chokepoint.",
        "auth.py: login(username, password) looks up the stored password hash for that "
        "username and compares it to hash_password(password) (SHA256), returning True only "
        "if they match (and False if the user does not exist). register(username, password) "
        "first checks whether the username is already taken; if not, it inserts a new row "
        "with the username and its hashed password. Passwords are hashed with SHA256 and "
        "never stored or compared in plain text.",
        "storage.py: save_strategy(username, strategy_data) inserts one row per saved "
        "strategy, storing name, type, tickers, start/end dates, params and kpis (tickers, "
        "params and kpis as JSONB). load_strategies(username) reads back every row for that "
        "user (oldest first), and for each one rebuilds the actual strategy object via "
        "strategies_config.STRATEGIES[row[\"type\"]] and the saved params (filtering out "
        "underscore-prefixed config keys like _cash/_risk_free, which are not constructor "
        "arguments), returning [] if the database cannot be reached. delete_strategy(username, "
        "name) removes one row by username and name, and update_kpis(username, name, kpis) "
        "overwrites just the kpis column after a re-run, leaving the rest of the saved "
        "configuration untouched.",
    ])
    add_para(doc, "Error handling: every Supabase call goes through run_query, which catches "
                  "exceptions, shows an error message, and returns None or an empty list so "
                  "callers fall back to their normal empty-state display. login and register "
                  "in auth.py treat that None as \"not found\" or \"taken\", so a database "
                  "outage fails safely. load_strategies in storage.py returns [] if the "
                  "database is unreachable, and defaults a saved strategy's missing "
                  "params/tickers to {}/[]. Not handled: if a saved strategy's type or params "
                  "no longer match the current strategies_config.STRATEGIES registry (for "
                  "example after a code change removes a parameter), rebuilding it with "
                  "strategy_cls(**constructor_args) would raise; this is not expected during "
                  "normal use since the registry and saved data are always written together.")


def write_ui(doc: Document) -> None:
    add_heading(doc, "Shared UI building blocks")
    add_para(doc, "theme.py holds every colour, font and spacing value as plain constants. "
                  "Both assets/style.css (via ui.inject_css) and the matplotlib charts (via "
                  "theme.apply_mpl_theme) read from these constants, so each colour is only "
                  "defined once. ui.py provides simple Streamlit wrappers used on every page: "
                  "page_header, kpi_card, verb_stepper, empty_state, next_step and data_table. "
                  "It also has two shared chart helpers: rebase_to_100 (indexes a series or "
                  "frame to 100, used for every equity-curve chart) and cml_chart (the "
                  "risk/return scatter with the Capital Market Line, used by Evaluate and "
                  "Compare). Error handling: cml_chart renders nothing if the benchmark has "
                  "zero or negative volatility, avoiding a division by zero in the Capital "
                  "Market Line slope.")
    add_para(doc, "strategies_config.py is the single registry the whole UI reads: each "
                  "strategy is listed once with its class, display label and the "
                  "strategy-specific parameters the Builder should ask for. Adding a strategy "
                  "is therefore a two-step job: write the class in core/strategy.py, then add "
                  "one registry entry. No page code needs to change.")
    add_para(doc, "prices.py adds a thin, UI-side st.cache_data layer around core/data.py (so "
                  "changing a dropdown does not re-download everything), plus two extra Yahoo "
                  "Finance lookups core/data.py does not provide: load_risk_free_rate and "
                  "load_info. Error handling: both catch any yfinance exception and fall back "
                  "to (0.04, \"\") and {} respectively, so a Yahoo outage degrades the page "
                  "instead of crashing it.")

    add_heading(doc, "Build -> Evaluate -> Compare", size=13)
    add_bullets(doc, [
        "Builder (pages/1_Builder.py): lets you set up a strategy (tickers, type and its "
        "registry-driven parameters, rebalancing frequency, an optional cash/risk-free "
        "overlay, and trading/management fees). _validate(...) collects every reason the "
        "strategy cannot be saved yet, and the page lists them together. On Save, it also "
        "loads prices for the selected tickers and date range, and adds an error naming any "
        "ticker that came back with no data (for example, a typo or a delisted symbol). On "
        "success, storage.save_strategy saves the configuration (not a backtest result).",
        "Evaluate (pages/2_Evaluate.py): rebuilds one saved strategy using the registry, "
        "runs it through Backtest alongside a benchmark (EqualWeight on the chosen "
        "benchmark ticker), applies the cash/risk-free overlay and fees to turn the gross "
        "curve into the net curve, and shows KPIs, factor models, charts, the CML and a "
        "forecast for that strategy.",
        "Compare (pages/3_Compare.py): does the same for two or more saved strategies, over "
        "the period they all share (the latest common start date to the earliest common "
        "end date). It ranks them in a Key Metrics table, a Costs table and a combined CML "
        "chart.",
    ])
    add_para(doc, "pages/0_Home.py is the entry point: it lists saved strategies with their "
                  "last-stored KPIs (so the page loads without re-running a backtest) and "
                  "links into Build/Evaluate/Compare. Its Update button re-runs Backtest with "
                  "fresh data up to yesterday and saves the refreshed KPIs.")
    add_para(doc, "Error handling: Backtest().run(...) in Home (Update), Evaluate and Compare "
                  "is wrapped in try/except. Home shows a toast and keeps the existing KPIs, "
                  "Evaluate shows a warning instead of the charts, and Compare skips the "
                  "failing strategy with a warning and continues with the others.")
    add_para(doc, "pages/4_Explore.py lets a user research tickers (price chart, CAGR/"
                  "volatility stats, fundamentals, correlation heatmap) before building a "
                  "strategy. For each selected ticker, it also shows a Forecast section "
                  "powered by core/forecast.py, with the model's historical test accuracy and "
                  "its predicted probability that tomorrow's return will be positive, with a "
                  "note that this is for educational purposes only. If a selected ticker comes "
                  "back with no price data (for example, a typo or a delisted symbol), the page "
                  "shows a warning naming it instead of silently showing fewer results.")
    add_para(doc, "pages/5_Glossary.py is a searchable, colour-coded reference explaining every "
                  "metric and strategy with its formula and an academic citation where "
                  "relevant (Markowitz, 1952; DeMiguel, Garlappi & Uppal, 2009; Jegadeesh & "
                  "Titman, 1993; Fama & French, 1993/2015).")


def write_collinearity(doc: Document) -> None:
    add_heading(doc, "Collinearity handling")
    add_para(doc, "Highly correlated inputs (near-identical tickers, or Fama-French factors "
                  "that move together) can make a covariance or design matrix singular. "
                  "Instead of crashing, MeanVariance and the factor regressions in "
                  "core/measures.py use np.linalg.pinv/lstsq (a pseudo-inverse / "
                  "least-squares solve) to still return a result, and the forecast model in "
                  "core/forecast.py relies on scikit-learn's default L2 (ridge) penalty to "
                  "keep correlated coefficients stable.")


def write_style(doc: Document) -> None:
    add_heading(doc, "Coding style")
    add_para(doc, "Small, single-purpose functions with type hints. Plain, readable loops "
                  "instead of clever one-liners, and short comments that explain why rather "
                  "than what. Every pages/*.py file follows the same layout: imports, then "
                  "session-state setup, then _-prefixed helper functions, then the page's "
                  "render code after a '# --- PAGE ---' marker. Streamlit code never appears "
                  "in core/.")


def write_ai_usage(doc: Document) -> None:
    add_heading(doc, "AI usage summary")
    add_para(doc, "Every source file carries its own \"AI Usage Declaration\" comment at the "
                  "top, describing what AI was used for in that specific file. This section "
                  "summarises that usage across the project.")
    add_bullets(doc, [
        "Claude was used throughout core/ and services/ to clean up, document and structure "
        "the code, and to explain specific concepts on request (for example, the "
        "@st.cache_resource vs @st.cache_data distinction, and why passwords should be "
        "hashed before storage).",
        "Claude Code was used to generate parts of the frontend code (app.py, the pages/, "
        "ui.py, theme.py, prices.py, strategies_config.py and tickers.py).",
        "ChatGPT was used to understand third-party tools and APIs: Supabase (the query "
        "builder, .execute(), JSONB columns) for the persistence layer, and Streamlit/"
        "matplotlib/yfinance concepts (session_state, caching, widgets, rcParams) for the "
        "frontend.",
        "The CSS in assets/style.css and the CSS assembled by ui.inject_css is AI-generated "
        "(Claude), because Streamlit's built-in theming is not flexible enough to restyle "
        "individual components.",
    ])
    add_para(doc, "In every case, the underlying logic and design decisions (what the app "
                  "does, how data flows, and how errors are handled) are our own; AI was used "
                  "for implementation help, clean-up, documentation and explaining unfamiliar "
                  "concepts.")


def build_document() -> None:
    doc = Document()
    #use a clean sans-serif default for body text
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    write_intro(doc)
    write_core(doc)
    write_services(doc)
    write_ui(doc)
    write_collinearity(doc)
    write_style(doc)
    write_ai_usage(doc)

    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    #the file is locked while it is open in Word or syncing in OneDrive - fail with a
    #clear instruction instead of a confusing traceback
    try:
        doc.save(OUTPUT_PATH)
    except PermissionError:
        print(f"could not write {OUTPUT_PATH} - it is probably open in Word or syncing "
              f"in OneDrive. close it and run 'python make_docs.py' again.")
        return
    print(f"wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
